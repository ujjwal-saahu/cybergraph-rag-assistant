from pathlib import Path
import pandas as pd
import pymupdf4llm
from docx import Document as DocxDocument

from app.utils.text_cleaner import clean_text


SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".txt",
    ".md",
    ".markdown",
    ".docx",
    ".csv",
}


def validate_file_extension(file_path: Path) -> None:
    """
    Validate uploaded file type.
    """

    extension = file_path.suffix.lower()

    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type: {extension}. "
            f"Supported types are: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )


def load_pdf_as_markdown(file_path: Path) -> str:
    """
    Convert PDF file into Markdown text using pymupdf4llm.
    """

    markdown_text = pymupdf4llm.to_markdown(str(file_path))
    return clean_text(markdown_text)


def load_text_file(file_path: Path) -> str:
    """
    Load TXT/Markdown file as plain text.
    """

    text = file_path.read_text(encoding="utf-8", errors="ignore")
    return clean_text(text)


def load_docx_as_markdown(file_path: Path) -> str:
    """
    Convert DOCX file into Markdown-like text.

    Paragraphs are preserved, and tables are converted into
    simple Markdown tables.
    """

    document = DocxDocument(str(file_path))

    blocks = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            blocks.append(text)

    for table_index, table in enumerate(document.tables, start=1):
        rows = []

        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append(cells)

        if not rows:
            continue

        blocks.append(f"\nTable {table_index}\n")

        header = rows[0]
        blocks.append("| " + " | ".join(header) + " |")
        blocks.append("| " + " | ".join(["---"] * len(header)) + " |")

        for row in rows[1:]:
            blocks.append("| " + " | ".join(row) + " |")

    return clean_text("\n\n".join(blocks))


def load_csv_as_markdown(file_path: Path) -> str:
    """
    Convert CSV file into Markdown text.

    For large CSV files, only the first 200 rows are included
    to avoid creating extremely large chunks.
    """

    try:
        df = pd.read_csv(file_path, encoding="utf-8")
    except UnicodeDecodeError:
        df = pd.read_csv(file_path, encoding="latin1")

    if df.empty:
        return ""

    row_count = len(df)
    column_names = list(df.columns)

    preview_df = df.head(200)

    markdown_table = preview_df.to_markdown(index=False)

    text = f"""
CSV File: {file_path.name}

Total rows: {row_count}
Total columns: {len(column_names)}
Columns: {", ".join(map(str, column_names))}

Data preview:

{markdown_table}
"""

    return clean_text(text)


def load_document(file_path: Path) -> str:
    """
    Load supported document and return clean Markdown/text.
    """

    validate_file_extension(file_path)

    extension = file_path.suffix.lower()

    if extension == ".pdf":
        return load_pdf_as_markdown(file_path)

    if extension in {".txt", ".md", ".markdown"}:
        return load_text_file(file_path)

    if extension == ".docx":
        return load_docx_as_markdown(file_path)

    if extension == ".csv":
        return load_csv_as_markdown(file_path)

    raise ValueError(f"Unsupported file type: {extension}")